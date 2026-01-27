#!/usr/bin/env python3
"""
Parse DOCX files from the Hedlin Family Journal.

This script extracts:
- Title/intro section
- Journal entries with dates
- Images embedded in documents
- Converts to structured JSON format for the build pipeline
"""

import argparse
import json
import re
import shutil
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Optional, List

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsmap
from PIL import Image
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn

console = Console()


# Date patterns to recognize journal entry dates
# Full format: "Saturday, July 23, 1983"
# Abbreviated: "Sunday, July 24" (year implied from context)
DATE_PATTERNS = [
    # Day of week + Month + Day + Year
    re.compile(
        r'^(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday),?\s+'
        r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+'
        r'(\d{1,2}),?\s+(\d{4})',
        re.IGNORECASE
    ),
    # Day of week + Month + Day (no year)
    re.compile(
        r'^(Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday),?\s+'
        r'(January|February|March|April|May|June|July|August|September|October|November|December)\s+'
        r'(\d{1,2})\b',
        re.IGNORECASE
    ),
    # Just Month + Day + Year
    re.compile(
        r'^(January|February|March|April|May|June|July|August|September|October|November|December)\s+'
        r'(\d{1,2}),?\s+(\d{4})',
        re.IGNORECASE
    ),
]

MONTH_MAP = {
    'january': 1, 'february': 2, 'march': 3, 'april': 4, 'may': 5, 'june': 6,
    'july': 7, 'august': 8, 'september': 9, 'october': 10, 'november': 11, 'december': 12
}


@dataclass
class JournalEntry:
    """A single journal entry."""
    date: str  # ISO date string (YYYY-MM-DD)
    date_display: str  # Original date text from document
    title: Optional[str] = None  # Optional title
    content: str = ""  # Entry content as markdown
    images: List[str] = None  # List of image filenames
    source_file: str = ""  # Source DOCX filename
    source_index: int = -1  # Paragraph index in source

    def __post_init__(self):
        if self.images is None:
            self.images = []


@dataclass
class ParsedDocument:
    """A parsed DOCX document containing journal entries."""
    title: str
    intro: str
    entries: List[JournalEntry]
    source_file: str
    parsed_at: str


def extract_date_from_text(text: str) -> Optional[tuple[datetime, str]]:
    """
    Extract a date from the given text.

    Returns (datetime, display_text) or None.
    """
    text = text.strip()

    for pattern in DATE_PATTERNS:
        match = pattern.match(text)
        if match:
            groups = match.groups()

            if len(groups) == 4:  # Full date with day, month, day, year
                day_name, month_name, day_num, year = groups
                month = MONTH_MAP[month_name.lower()]
                try:
                    date_obj = datetime(int(year), month, int(day_num))
                    return date_obj, match.group(0)
                except ValueError:
                    continue

            elif len(groups) == 3:
                # Could be (day, month, day) abbreviated, or (month, day, year)
                if groups[0] in MONTH_MAP:  # (month, day, year)
                    month_name, day_num, year = groups
                    month = MONTH_MAP[month_name.lower()]
                    try:
                        date_obj = datetime(int(year), month, int(day_num))
                        return date_obj, match.group(0)
                    except ValueError:
                        continue
                else:  # (day, month, day) - abbreviated, no year
                    # We'll need context to determine the year
                    return None, match.group(0)

    return None, None


def extract_images(doc: Document, output_dir: Path, source_name: str) -> dict[int, list[str]]:
    """
    Extract images from the document.

    Returns a mapping of paragraph index to list of image filenames.
    """
    image_map = {}
    images_dir = output_dir / "content" / "images"
    images_dir.mkdir(parents=True, exist_ok=True)

    # Get image relationships
    image_rels = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.target_ref:
            image_rels[rel_id] = rel

    if not image_rels:
        console.print("[dim]No images found in document[/dim]")
        return image_map

    # Extract images and track which paragraphs they appear in
    image_counter = 0
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_images = []

        # Check for images in this paragraph
        for run in paragraph.runs:
            for blip in run._element.xpath('.//a:blip'):
                embed_id = blip.get(nsmap('a')('embed'))
                if embed_id and embed_id in image_rels:
                    image = image_rels[embed_id].target_part.blob
                    image_counter += 1

                    # Detect format and save
                    img = Image.open(io.BytesIO(image))
                    ext = img.format.lower() if img.format else 'png'
                    filename = f"{source_name}_{image_counter:04d}.{ext}"
                    filepath = images_dir / filename

                    with open(filepath, 'wb') as f:
                        f.write(image)

                    # Store relative path
                    para_images.append(f"/content/images/{filename}")
                    console.print(f"  [dim]Extracted image: {filename}[/dim]")

        if para_images:
            image_map[para_idx] = para_images

    return image_map


import io


def extract_year_from_filename(filename: str) -> Optional[int]:
    """Extract a 4-digit year from a filename (e.g., 'Hedlin Family Journal 1984.docx' -> 1984)."""
    match = re.search(r'\b(19\d{2}|20\d{2})\b', filename)
    if match:
        return int(match.group(1))
    return None


def parse_docx_file(
    docx_path: Path,
    output_dir: Path,
    default_year: Optional[int] = None
) -> ParsedDocument:
    """
    Parse a single DOCX file into structured entries.

    Args:
        docx_path: Path to the DOCX file
        output_dir: Directory to write extracted content
        default_year: Default year for abbreviated dates (e.g., "Sunday, July 24")

    Returns:
        ParsedDocument with title, intro, and entries
    """
    console.print(f"\n[bold cyan]Parsing:[/bold cyan] {docx_path.name}")

    doc = Document(docx_path)
    source_name = docx_path.stem

    # Extract images first (we need to know which paragraphs have images)
    image_map = extract_images(doc, output_dir, source_name)

    entries = []
    current_entry: Optional[JournalEntry] = None
    current_content = []
    intro_lines = []
    title = ""

    # First pass: find the title (usually first non-empty line)
    for para in doc.paragraphs[:20]:
        text = para.text.strip()
        if text and not title:
            title = text
            break

    # Set default year: 1) from filename, 2) from first full date, 3) use 1900 as safe default
    if default_year is None:
        default_year = extract_year_from_filename(docx_path.name)
        if default_year:
            console.print(f"[dim]Using year from filename: {default_year}[/dim]")

    if default_year is None:
        for para in doc.paragraphs:
            date_obj, _ = extract_date_from_text(para.text)
            if date_obj:
                default_year = date_obj.year
                console.print(f"[dim]Detected default year from content: {default_year}[/dim]")
                break

    last_year = default_year or 1900

    # Second pass: extract entries
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Check if this is a date header
        date_obj, date_display = extract_date_from_text(text)

        if date_display:
            # Save previous entry if exists
            if current_entry:
                current_entry.content = '\n\n'.join(current_content)
                entries.append(current_entry)

            # Handle abbreviated dates (no year)
            if date_obj:
                entry_date = date_obj
                last_year = entry_date.year
            else:
                # Use last known year
                month_str = re.search(r'(January|February|March|April|May|June|July|August|September|October|November|December)', text, re.IGNORECASE)
                if month_str:
                    month = MONTH_MAP[month_str.group(0).lower()]
                    day_str = re.search(r'\b(\d{1,2})\b', text)
                    if day_str:
                        try:
                            entry_date = datetime(last_year, month, int(day_str.group(0)))
                        except ValueError:
                            entry_date = datetime.now()
                    else:
                        entry_date = datetime.now()
                else:
                    entry_date = datetime.now()

            # Create new entry
            current_entry = JournalEntry(
                date=entry_date.strftime('%Y-%m-%d'),
                date_display=date_display,
                source_file=docx_path.name,
                source_index=para_idx,
                images=image_map.get(para_idx, [])
            )
            current_content = []

        elif current_entry:
            # Add content to current entry
            current_content.append(text)
            # Check if this paragraph has images
            if para_idx in image_map:
                if current_entry.images is None:
                    current_entry.images = []
                current_entry.images.extend(image_map[para_idx])

        else:
            # Before first entry - collect intro
            intro_lines.append(text)

    # Save last entry
    if current_entry:
        current_entry.content = '\n\n'.join(current_content)
        entries.append(current_entry)

    # Build intro from lines before first entry
    intro = '\n\n'.join(intro_lines[:10])  # Limit intro length

    console.print(f"[green]  Found {len(entries)} entries[/green]")

    return ParsedDocument(
        title=title,
        intro=intro,
        entries=entries,
        source_file=docx_path.name,
        parsed_at=datetime.now().isoformat()
    )


def find_docx_files(docs_dir: Path) -> List[Path]:
    """Find all DOCX files in the given directory."""
    files = list(docs_dir.glob("*.docx"))
    files.extend(docs_dir.glob("*.DOCX"))
    return sorted(set(files))


def merge_entries(all_entries: List[JournalEntry]) -> List[JournalEntry]:
    """
    Merge entries from multiple documents, deduplicating by date.
    Later documents override earlier ones for the same date.
    """
    entry_map = {}

    for entry in all_entries:
        key = (entry.date, entry.source_file, entry.source_index)
        if key not in entry_map:
            entry_map[key] = entry

    # Sort by date
    sorted_entries = sorted(entry_map.values(), key=lambda e: e.date)

    return sorted_entries


def main():
    parser = argparse.ArgumentParser(
        description="Parse Hedlin Family Journal DOCX files"
    )
    parser.add_argument(
        "--input", "-i",
        type=Path,
        default=Path("docs"),
        help="Directory containing DOCX files (default: docs/)"
    )
    parser.add_argument(
        "--output", "-o",
        type=Path,
        default=Path("data"),
        help="Directory for output JSON files (default: data/)"
    )
    parser.add_argument(
        "--year", "-y",
        type=int,
        default=None,
        help="Default year for abbreviated dates"
    )
    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Verbose output"
    )

    args = parser.parse_args()

    # Create output directory
    args.output.mkdir(parents=True, exist_ok=True)
    project_root = Path.cwd()

    # Find DOCX files
    docx_files = find_docx_files(args.input)

    if not docx_files:
        console.print("[red]No DOCX files found in {args.input}[/red]")
        return 1

    console.print(f"[bold]Found {len(docx_files)} DOCX file(s)[/bold]")

    # Parse all documents
    all_parsed_docs = []
    all_entries = []

    with Progress(
        SpinnerColumn(),
        TextColumn("[progress.description]{task.description}"),
        console=console
    ) as progress:
        task = progress.add_task("Parsing documents...", total=len(docx_files))

        for docx_file in docx_files:
            parsed = parse_docx_file(docx_file, project_root, args.year)
            all_parsed_docs.append(parsed)
            all_entries.extend(parsed.entries)
            progress.update(task, advance=1)

    # Write combined JSON output
    output_file = args.output / "journal_entries.json"

    combined_data = {
        "title": "Hedlin Family Journal",
        "parsed_at": datetime.now().isoformat(),
        "documents": [
            {
                "title": doc.title,
                "source_file": doc.source_file,
                "intro": doc.intro,
                "entry_count": len(doc.entries)
            }
            for doc in all_parsed_docs
        ],
        "entries": [
            {
                "date": entry.date,
                "date_display": entry.date_display,
                "title": entry.title,
                "content": entry.content,
                "images": entry.images,
                "source_file": entry.source_file
            }
            for entry in sorted(all_entries, key=lambda e: e.date)
        ]
    }

    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(combined_data, f, indent=2, ensure_ascii=False)

    console.print(f"\n[green]Output written to:[/green] {output_file}")
    console.print(f"[green]Total entries: {len(all_entries)}[/green]")

    return 0


if __name__ == "__main__":
    exit(main())
